import os
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
from google.adk.sessions import InMemorySessionService
from google.adk.runners import Runner
import json
import asyncio
from core.config import file_io_lock, logger, ephemeral_audit_logs

file_io_lock = asyncio.Lock()

# ==========================================
# 1. GOOGLE ADK UTILITIES
# ==========================================

def _make_runner(agent):
    return Runner(agent=agent, app_name="askFINN", session_service=InMemorySessionService())

def extract_text_from_events(events):
    """
    Robustly extracts the final text response from a list of Agent events.
    Handles 'content.parts', 'parts', and 'output' attributes.
    """
    final_text = ""
    
    print(f"🔍 DEBUG: Scanning {len(events)} events for output...")

    for i, event in enumerate(events):
        # 1. Check for 'content.parts' (The structure seen in your logs)
        if hasattr(event, "content") and hasattr(event.content, "parts"):
            for part in event.content.parts:
                if hasattr(part, "text") and part.text:
                    final_text = part.text
        
        # 2. Check for direct '.parts' (Older Google GenAI structure)
        elif hasattr(event, "parts"):
            for part in event.parts:
                if hasattr(part, "text") and part.text:
                    final_text = part.text

        # 3. Check standard '.text'
        elif hasattr(event, "text") and event.text:
            final_text = event.text
            
        # 4. Check for Tool Output (e.g. Chart JSON directly in event)
        elif hasattr(event, "output") and event.output:
            final_text = str(event.output)

    # Decision Time
    if final_text:
        print(f"✅ Found Output: {final_text[:100]}...") 
        return final_text


# ==========================================
# 2. DOCUMENT GENERATOR (The Payoff!)
# ==========================================

def generate_rdec_docx(aif_state: dict, session_id: str):
    """
    Generates a technical-only RDEC document. 
    Removes Company Details, Project Summaries, and Financial tables.
    """
    print(f"📝 Generating Technical-Only Word Document for session {session_id}...")
    
    doc = Document()
    
    # --- Title ---
    title = doc.add_heading('RDEC Technical Narrative Report', 0)
    title.alignment = 1 
    doc.add_paragraph(f"Report Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n").alignment = 1

    # --- Section 1: Project Narratives ---
    # This is now the primary focus of the document
    doc.add_heading('1. Project Technical Narratives', level=1)
    narratives = aif_state.get('project_narratives', [])
    
    if not narratives:
        doc.add_paragraph("No technical narratives were generated for this session.")
    else:
        for idx, proj in enumerate(narratives, 1):
            doc.add_heading(f"Project: {proj.get('project_name', 'Unnamed Project')}", level=2)
            
            # Competent Professional is critical to the narrative
            doc.add_paragraph(f"Lead Competent Professional: {proj.get('competent_professional', 'Not Provided')}")
            
            # Technical Body
            doc.add_heading('Advance in Science or Technology', level=3)
            doc.add_paragraph(proj.get('advance_sought', 'Not Provided'))
            
            doc.add_heading('Technological Uncertainties', level=3)
            doc.add_paragraph(proj.get('scientific_uncertainties', 'Not Provided'))
            
            doc.add_heading('Resolution of Uncertainties', level=3)
            doc.add_paragraph(proj.get('why_unresolvable_by_professional', 'Not Provided'))
            
            doc.add_heading('Technical Activities Undertaken', level=3)
            doc.add_paragraph(proj.get('activities_undertaken', 'Not Provided'))
            
            doc.add_heading('Project Outcomes', level=3)
            doc.add_paragraph(proj.get('outcomes', 'Not Provided'))
            
            # Page break between projects if there are multiple
            if len(narratives) > 1 and idx < len(narratives):
                doc.add_page_break()

    # --- Section 2: Mandatory Compliance Flags ---
    doc.add_heading('2. Mandatory Compliance Flags', level=1)
    comp = aif_state.get('compliance', {})
    
    def yes_no(val):
        if val is None: return "Not Answered"
        return "Yes" if val else "No"

    # Using the standardized keys we aligned earlier
    doc.add_paragraph(f"Overseas R&D involved? {yes_no(comp.get('overseas_rnd'))}")
    doc.add_paragraph(f"AI or Machine Learning utilized? {yes_no(comp.get('ai_used'))}")
    doc.add_paragraph(f"Quantum Technologies involved? {yes_no(comp.get('quantum_used'))}")
    
    # --- Save Logic ---
    export_dir = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'data', 'exports')
    os.makedirs(export_dir, exist_ok=True)
    
    file_path = os.path.join(export_dir, f"AIF_Technical_{session_id}.docx")
    doc.save(file_path)
    
    print(f"✅ Technical document saved: {file_path}")
    return file_path

def deep_merge(dict1, dict2):
    """Recursively merges dict2 into dict1 without dropping existing keys."""
    for key, value in dict2.items():
        if isinstance(value, dict) and key in dict1 and isinstance(dict1[key], dict):
            deep_merge(dict1[key], value)
        else:
            dict1[key] = value
    return dict1

def determine_next_field(aif_state: dict) -> str:
    """Properly scans the state dictionary for missing or weak fields."""
    
    # 1. Check the Project Narratives Array
    narratives = aif_state.get("project_narratives", [])
    if not narratives:
        return "project_name" # Fallback if array is completely empty
        
    proj = narratives[0]
    
    # Define the exact sequential order we want to ask questions in
    narrative_keys = [
        "project_name", 
        "competent_professional", 
        "advance_sought", 
        "scientific_uncertainties", 
        "why_unresolvable_by_professional", 
        "activities_undertaken", 
        "outcomes"
    ]
    
    for key in narrative_keys:
        val = proj.get(key)
        # Check if it's missing entirely
        if val is None or val == "":
            return key
        # Check if Python flagged it as a weak draft
        if isinstance(val, str) and "[WEAK_DRAFT]" in val:
            return key
            
    # 2. Check the Compliance Fields (Aligned with DOCX generator)
    compliance = aif_state.get("compliance", {})
    # 🐛 FIXED: Keys now perfectly match the generate_rdec_docx schema
    compliance_keys = ["overseas_rnd", "ai_used", "quantum_used"] 
    
    for key in compliance_keys:
        val = compliance.get(key)
        # For booleans, we strictly check for None (since False is a valid answer)
        if val is None:
            return key
            
    # If it survives all those checks, it is TRULY complete.
    return "Complete"

async def append_to_master_log(SUB_dir, SAVED_dir, session_id: str, actor: str, event_type: str, details: str):
    # 1. Always store in the temporary memory buffer first
    if session_id not in ephemeral_audit_logs:
        ephemeral_audit_logs[session_id] = []
        
    new_entry = {
        "timestamp": datetime.now().isoformat(),
        "actor": actor,
        "event_type": event_type,
        "details": details
    }
    ephemeral_audit_logs[session_id].append(new_entry)

    # 2. Try to find if the user has ALREADY saved a file
    saved_path = os.path.join(SAVED_dir, f"{session_id}.json")
    sub_path = os.path.join(SUB_dir, f"{session_id}.json")
    
    target_path = sub_path if os.path.exists(sub_path) else (saved_path if os.path.exists(saved_path) else None)

    # 3. If a file exists, flush the memory to the disk!
    if target_path:
        async with file_io_lock:
            with open(target_path, "r") as f:
                data = json.load(f)
            
            if "audit_summary" not in data or data["audit_summary"] is None: 
                data["audit_summary"] = {}
            if "detailed_log" not in data["audit_summary"]: 
                data["audit_summary"]["detailed_log"] = []
            
            # Dump all pending memory logs into the file
            data["audit_summary"]["detailed_log"].extend(ephemeral_audit_logs[session_id])
            
            with open(target_path, "w") as f:
                json.dump(data, f, indent=2)
            
            # Clear the temporary memory now that it's safely on disk
            ephemeral_audit_logs[session_id] = []